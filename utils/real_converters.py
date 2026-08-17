"""
ModelFlow Real File Converters & Data Processing Engine
Handles actual binary file conversions for PDF, Word, Images, Markdown, and Developer Formats
"""

import io
import os
import json
import csv
import re
import xml.etree.ElementTree as ET
from PIL import Image

from utils.tools_engine import (
    convert_json_to_yaml, convert_yaml_to_json, convert_csv_to_json,
    convert_json_to_xml, convert_xml_to_json
)

# Import backend document libraries
import pypdf
import docx
import markdown
import yaml
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
from reportlab.lib import colors


def images_to_pdf_bytes(image_files_data):
    """
    Converts multiple image files (JPG, PNG, WebP, BMP) into a single PDF document.
    """
    pdf_buffer = io.BytesIO()
    pil_images = []

    for img_bytes in image_files_data:
        try:
            im = Image.open(io.BytesIO(img_bytes))
            if im.mode in ("RGBA", "P"):
                im = im.convert("RGB")
            pil_images.append(im)
        except Exception:
            continue

    if not pil_images:
        raise ValueError("No valid image files provided for PDF conversion.")

    first_image = pil_images[0]
    rest_images = pil_images[1:]
    first_image.save(pdf_buffer, format="PDF", save_all=True, append_images=rest_images)
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()


def image_format_convert(image_bytes, target_format="PNG"):
    """
    Converts image bytes between JPG, PNG, WebP, BMP, TIFF.
    """
    im = Image.open(io.BytesIO(image_bytes))
    target_fmt = target_format.upper()
    if target_fmt == "JPG":
        target_fmt = "JPEG"
    
    if target_fmt in ("JPEG", "JPG") and im.mode in ("RGBA", "P"):
        im = im.convert("RGB")

    out_buffer = io.BytesIO()
    im.save(out_buffer, format=target_fmt)
    out_buffer.seek(0)
    return out_buffer.getvalue()


def text_or_md_to_pdf_bytes(text_content, is_markdown=True):
    """
    Converts Markdown or plain text content into a styled PDF document using ReportLab.
    """
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    story = []

    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#1e293b')
    )
    h1_style = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        textColor=colors.HexColor('#4f46e5'),
        spaceAfter=10
    )

    if is_markdown:
        html_content = markdown.markdown(text_content)
        # Parse basic headings and paragraphs from HTML
        lines = html_content.splitlines()
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
            if line_str.startswith("<h1>"):
                clean = re.sub(r'<[^>]+>', '', line_str)
                story.append(Paragraph(clean, h1_style))
                story.append(Spacer(1, 10))
            elif line_str.startswith("<h2>") or line_str.startswith("<h3>"):
                clean = re.sub(r'<[^>]+>', '', line_str)
                story.append(Paragraph(f"<b>{clean}</b>", h1_style))
                story.append(Spacer(1, 8))
            else:
                clean = re.sub(r'<[^>]+>', '', line_str)
                story.append(Paragraph(clean, body_style))
                story.append(Spacer(1, 6))
    else:
        for para in text_content.split("\n\n"):
            clean = para.strip().replace("\n", "<br/>")
            if clean:
                story.append(Paragraph(clean, body_style))
                story.append(Spacer(1, 8))

    doc.build(story)
    pdf_buffer.seek(0)
    return pdf_buffer.getvalue()


def word_docx_to_pdf_bytes(docx_bytes):
    """
    Parses Microsoft Word DOCX bytes and builds a PDF file.
    """
    doc_in = docx.Document(io.BytesIO(docx_bytes))
    full_text = []
    for p in doc_in.paragraphs:
        if p.text.strip():
            full_text.append(p.text.strip())
    
    text_content = "\n\n".join(full_text)
    return text_or_md_to_pdf_bytes(text_content, is_markdown=False)


def pdf_to_word_bytes(pdf_bytes):
    """
    Extracts text from PDF bytes and builds an editable Word DOCX file.
    """
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    doc_out = docx.Document()
    doc_out.add_heading('Extracted Document Content', level=1)

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        doc_out.add_heading(f"Page {i+1}", level=2)
        for line in page_text.splitlines():
            line_str = line.strip()
            if line_str:
                doc_out.add_paragraph(line_str)

    docx_buffer = io.BytesIO()
    doc_out.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer.getvalue()


def pdf_to_text_string(pdf_bytes):
    """
    Extracts plain text from a PDF file.
    """
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    extracted = []
    for i, page in enumerate(reader.pages):
        extracted.append(f"--- Page {i+1} ---")
        extracted.append(page.extract_text() or "")
    return "\n\n".join(extracted)


def pdf_to_markdown_string(pdf_bytes):
    """
    Extracts text from PDF and converts into structured Markdown syntax.
    """
    raw_text = pdf_to_text_string(pdf_bytes)
    lines = raw_text.splitlines()
    md_lines = []
    for line in lines:
        l = line.strip()
        if not l:
            continue
        if l.startswith("--- Page"):
            md_lines.append(f"## {l}")
        elif len(l) < 40 and not l.endswith("."):
            md_lines.append(f"### {l}")
        else:
            md_lines.append(l)
    return "\n\n".join(md_lines)


def pdf_merge_bytes(pdf_files_list):
    """
    Merges multiple PDF byte arrays into a single unified PDF file.
    """
    merger = pypdf.PdfMerger()
    for pdf_b in pdf_files_list:
        merger.append(io.BytesIO(pdf_b))

    merged_buffer = io.BytesIO()
    merger.write(merged_buffer)
    merger.close()
    merged_buffer.seek(0)
    return merged_buffer.getvalue()


def pdf_split_bytes(pdf_bytes, page_range_str="1-2"):
    """
    Splits specific page ranges from a PDF document.
    """
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    total_pages = len(reader.pages)

    target_pages = []
    try:
        if "-" in page_range_str:
            start_p, end_p = page_range_str.split("-")
            start_i = max(0, int(start_p.strip()) - 1)
            end_i = min(total_pages, int(end_p.strip()))
            target_pages = list(range(start_i, end_i))
        elif "," in page_range_str:
            target_pages = [int(p.strip()) - 1 for p in page_range_str.split(",") if p.strip().isdigit()]
        else:
            target_pages = [int(page_range_str.strip()) - 1]
    except Exception:
        target_pages = list(range(min(2, total_pages)))

    for p_idx in target_pages:
        if 0 <= p_idx < total_pages:
            writer.add_page(reader.pages[p_idx])

    out_buf = io.BytesIO()
    writer.write(out_buf)
    out_buf.seek(0)
    return out_buf.getvalue()


def pdf_rotate_bytes(pdf_bytes, angle=90):
    """
    Rotates all pages in a PDF document by 90, 180, or 270 degrees.
    """
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    for page in reader.pages:
        page.rotate(angle)
        writer.add_page(page)

    out_buf = io.BytesIO()
    writer.write(out_buf)
    out_buf.seek(0)
    return out_buf.getvalue()


def pdf_protect_bytes(pdf_bytes, password):
    """
    Encrypts a PDF document with a user password.
    """
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    writer = pypdf.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    
    writer.encrypt(password)
    out_buf = io.BytesIO()
    writer.write(out_buf)
    out_buf.seek(0)
    return out_buf.getvalue()


def pdf_remove_password_bytes(pdf_bytes, password):
    """
    Decrypts a PDF document using a user password.
    """
    reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
    if reader.is_encrypted:
        reader.decrypt(password)
    
    writer = pypdf.PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    out_buf = io.BytesIO()
    writer.write(out_buf)
    out_buf.seek(0)
    return out_buf.getvalue()
